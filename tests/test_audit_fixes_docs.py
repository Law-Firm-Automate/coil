"""Regression tests for three audit defects.

1. docs.json "Document automation from DOCX and HTML templates" (partial, medium):
   plain {FIELD} merge fields are detected and offered on the generate form but never substituted.
2. docs.json "Citation check that flags citations which do not resolve" (partial, medium):
   a CourtListener status 300 (ambiguous) is labelled "not found, verify before filing".
3. clients.json "Draft client update email from notes, work, tasks and dates" (partial, low):
   the template fallback copies matter notes verbatim, so an AI summary note saved by Coil itself
   leaks fee and hours language into a client-facing draft.

Own SQLite file and own UPLOAD_DIR/PDF_DIR. No network: the CourtListener transport is monkeypatched
and the AI keys are blanked so app.llm raises LLMUnavailable.
Run: .venv/bin/python -m pytest tests/test_audit_fixes_docs.py -q
"""
import io
import os
import re
import shutil
import subprocess
import sys
import zipfile

import pytest

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, ROOT)
from tests.helpers import login  # noqa: E402

DB_PATH = os.path.join(ROOT, "data", "test_audit_fixes_docs.db")
DB_URI = f"sqlite:///{DB_PATH}"
UPLOAD_DIR = os.path.join(ROOT, "data", "uploads", "test_audit_fixes_docs")
PDF_DIR = os.path.join(ROOT, "data", "pdf", "test_audit_fixes_docs")


@pytest.fixture(scope="module")
def app():
    if os.path.exists(DB_PATH):
        os.remove(DB_PATH)
    shutil.rmtree(UPLOAD_DIR, ignore_errors=True)
    shutil.rmtree(PDF_DIR, ignore_errors=True)
    env = dict(os.environ, DATABASE_URL=DB_URI)
    out = subprocess.run([sys.executable, os.path.join(ROOT, "seed.py")], env=env, cwd=ROOT,
                         capture_output=True, text=True)
    assert out.returncode == 0, out.stderr
    from app import create_app
    a = create_app({"SQLALCHEMY_DATABASE_URI": DB_URI, "UPLOAD_DIR": UPLOAD_DIR, "PDF_DIR": PDF_DIR,
                    "TESTING": True, "SMTP_HOST": "", "OPENROUTER_API_KEY": "", "ANTHROPIC_API_KEY": "",
                    "COURTLISTENER_TOKEN": ""})
    yield a


@pytest.fixture
def owner(app):
    c = app.test_client()
    tok = login(c)
    return c, tok


@pytest.fixture(autouse=True)
def no_ai_keys(monkeypatch):
    """app.llm reads keys at call time, so a key in the developer's shell must not leak in."""
    for k in ("OPENROUTER_API_KEY", "ANTHROPIC_API_KEY", "LLM_ENABLED", "LLM_DAILY_CAP", "AI_DAILY_CAP_CENTS",
              "COURTLISTENER_TOKEN"):
        monkeypatch.delenv(k, raising=False)


def _models():
    from app.extensions import db
    from app import models
    return db, models


def _matter(app, number="M-1001"):
    from app.models import Matter
    with app.app_context():
        m = Matter.query.filter_by(number=number).first()
        return m.id, m.client.display_name


def _docx_text(data):
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    return re.sub(r"<[^>]+>", "", xml)


def _pdf_text(data):
    from pypdf import PdfReader
    return "\n".join((p.extract_text() or "") for p in PdfReader(io.BytesIO(data)).pages)


# ---------------------------------------------------------------------------
# 1. plain {FIELD} merge fields
# ---------------------------------------------------------------------------
def _split_field_docx():
    """A .docx whose {CLIENT_NAME} field is deliberately split across two Word runs."""
    from docx import Document as DocxDocument
    d = DocxDocument()
    d.add_paragraph("Dear {{ client_name }},")
    p = d.add_paragraph()
    p.add_run("Plain style field: {CLIENT")
    p.add_run("_NAME} is the payer.")
    d.add_paragraph("Claim {CLAIM_NUMBER} on file {{ matter_number }}.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_plain_field_docx_split_across_runs_is_filled(app, owner):
    c, tok = owner
    mid, client_name = _matter(app, "M-1001")
    r = c.post("/doctemplates/new",
               data={"_csrf": tok, "name": "Plain field letter", "kind": "docx", "practice_area": "",
                     "description": "", "is_active": "1", "file": (io.BytesIO(_split_field_docx()), "plain.docx")},
               content_type="multipart/form-data")
    assert r.status_code == 302, r.data[:400]
    tid = int(re.search(r"/doctemplates/(\d+)/edit", r.headers["Location"]).group(1))
    from app.models import DocTemplate, Document
    with app.app_context():
        assert "CLIENT_NAME" in DocTemplate.query.get(tid).fields

    # the generate form offers the plain field
    r = c.get(f"/doctemplates/{tid}/generate?matter_id={mid}")
    assert r.status_code == 200 and 'name="f_CLIENT_NAME"' in r.data.decode()

    r = c.post(f"/doctemplates/{tid}/generate",
               data={"_csrf": tok, "matter_id": str(mid), "f_client_name": client_name,
                     "f_matter_number": "M-1001", "f_CLIENT_NAME": "Acme Holdings LLC",
                     "f_CLAIM_NUMBER": "CLM-77"})
    assert r.status_code == 302, r.data[:400]
    with app.app_context():
        doc = Document.query.filter_by(matter_id=mid, template_id=tid).first()
        assert doc is not None
        text = _docx_text(open(os.path.join(UPLOAD_DIR, doc.path), "rb").read())
    assert "Acme Holdings LLC" in text, "plain {FIELD} split across runs was not substituted"
    assert "CLM-77" in text
    assert "{CLIENT_NAME}" not in text and "{CLAIM_NUMBER}" not in text
    assert "{{" not in text and "{" not in text


def test_plain_field_html_template_is_filled(app, owner):
    c, tok = owner
    mid, client_name = _matter(app, "M-1001")
    body = ("<style>p { color: #111; margin: 0 }</style>"
            "<p>Client of record: {CLIENT_NAME}</p><p>File {{ matter_number }}</p>")
    r = c.post("/doctemplates/new", data={"_csrf": tok, "name": "Plain field HTML", "kind": "html",
                                          "practice_area": "", "description": "", "is_active": "1",
                                          "body_html": body})
    assert r.status_code == 302, r.data[:400]
    tid = int(re.search(r"/doctemplates/(\d+)/edit", r.headers["Location"]).group(1))
    r = c.post(f"/doctemplates/{tid}/generate",
               data={"_csrf": tok, "matter_id": str(mid), "f_matter_number": "M-1001",
                     "f_CLIENT_NAME": "Acme Holdings LLC"})
    assert r.status_code == 302, r.data[:400]
    from app.models import Document
    with app.app_context():
        doc = Document.query.filter_by(matter_id=mid, template_id=tid).first()
        assert doc is not None and doc.mime == "application/pdf"
        pdf = open(os.path.join(UPLOAD_DIR, doc.path), "rb").read()
    text = _pdf_text(pdf)
    assert "Acme Holdings LLC" in text, "plain {FIELD} in an HTML template was not substituted"
    assert "{CLIENT_NAME}" not in text


def test_plain_field_substitution_leaves_css_and_unknown_fields_alone():
    """Braces that are not a supplied merge field must survive untouched."""
    from app.blueprints.doctemplates import fill_plain_fields
    html = "<style>p { color: red } .a{margin:0}</style><p>{CLIENT_NAME} and {OTHER_FIELD}</p>"
    out = fill_plain_fields(html, {"CLIENT_NAME": "Acme & Sons"})
    assert "p { color: red }" in out and ".a{margin:0}" in out
    assert "{OTHER_FIELD}" in out
    assert "Acme &amp; Sons" in out and "{CLIENT_NAME}" not in out


# ---------------------------------------------------------------------------
# 2. ambiguous citations
# ---------------------------------------------------------------------------
AMBIGUOUS_JSON = [{
    "citation": "1 X. 1", "normalized_citations": ["1 X. 1"], "status": 300, "error_message": "",
    "start_index": 4, "end_index": 11,
    "clusters": [
        {"id": 1, "case_name": "A v. B", "date_filed": "2000-01-01", "absolute_url": "/opinion/1/a/", "citations": []},
        {"id": 2, "case_name": "C v. D", "date_filed": "2001-01-01", "absolute_url": "/opinion/2/c/", "citations": []},
    ],
}, {
    "citation": "9 Z. 9", "normalized_citations": ["9 Z. 9"], "status": 404, "error_message": "",
    "start_index": 20, "end_index": 27, "clusters": [],
}]


def test_status_300_is_ambiguous_not_not_found(app, owner, monkeypatch):
    c, tok = owner
    db, M = _models()
    mid, _ = _matter(app, "M-1001")
    from app.blueprints import _courtlistener as cl
    cl.clear_cache()
    monkeypatch.setattr(cl, "_post", lambda path, data=None: {"ok": True, "data": AMBIGUOUS_JSON})

    r = c.post("/research/cite-check",
               data={"_csrf": tok, "text": "See 1 X. 1 and also 9 Z. 9.", "matter_id": str(mid)})
    assert r.status_code == 200
    html = r.data.decode()

    # the ambiguous row must not claim the citation was not found, and must list its candidates
    row = html.split("<strong>1 X. 1</strong>")[1].split("</tr>")[0]
    assert "not found" not in row, "status 300 is still labelled not found"
    assert "possible matches" in row
    assert "A v. B" in row and "C v. D" in row, "candidate clusters are not listed"
    assert "verify" in row.lower()
    # the genuine 404 keeps the not-found wording
    row404 = html.split("<strong>9 Z. 9</strong>")[1].split("</tr>")[0]
    assert "not found, verify before filing" in row404
    # header counts separate the three outcomes
    assert "1 ambiguous" in html and "1 not found" in html

    with app.app_context():
        n = M.Note.query.filter_by(matter_id=mid).order_by(M.Note.id.desc()).first()
        assert n is not None and "Citation check (CourtListener)" in n.body
        line = [x for x in n.body.splitlines() if x.startswith("- 1 X. 1")][0]
        assert "not found" not in line, "the matter note still records the ambiguous cite as not found"
        assert "2 possible matches" in line and "A v. B" in line
        assert "- 9 Z. 9: not found, verify before filing" in n.body


# ---------------------------------------------------------------------------
# 3. client update email must not carry attorney work product
# ---------------------------------------------------------------------------
def test_client_update_draft_excludes_internal_ai_summary(app, owner):
    c, tok = owner
    db, M = _models()
    mid, _ = _matter(app, "M-1002")
    with app.app_context():
        db.session.add(M.Note(matter_id=mid, user_id=1, body="Filed the motion to compel."))
        db.session.commit()

    # Coil's own AI summary, saved to the matter by /ai/matter/<id>/summary/save
    r = c.post(f"/ai/matter/{mid}/summary/save",
               data={"_csrf": tok,
                     "summary": "Client is elderly and may lack capacity; get a doctor letter before signing.",
                     "open_items": "Bill 2 unbilled hours\nChase the unpaid invoice"})
    assert r.status_code == 302
    with app.app_context():
        n = M.Note.query.filter_by(matter_id=mid).order_by(M.Note.id.desc()).first()
        assert n.body.lstrip().lower().startswith("[internal]"), \
            "an app-generated work-product note is not marked internal"

    # no AI keys, so the deterministic template fallback runs
    r = c.post(f"/ai/matter/{mid}/update-email", data={"_csrf": tok})
    assert r.status_code == 200
    html = r.data.decode()
    draft = html.split('name="body"')[1].split("</textarea>")[0]
    assert "may lack capacity" not in draft, "attorney work product leaked into the client draft"
    assert "unbilled hours" not in draft and "unpaid invoice" not in draft
    assert "AI summary" not in draft
    # ordinary notes still reach the client draft
    assert "Filed the motion to compel." in draft


def test_update_facts_drops_notes_that_talk_about_money(app):
    db, M = _models()
    mid, _ = _matter(app, "M-1001")
    with app.app_context():
        from app.blueprints.ai import update_facts
        db.session.add(M.Note(matter_id=mid, user_id=1, body="Reviewed the deposition transcript."))
        db.session.add(M.Note(matter_id=mid, user_id=1, body="Client still owes $2,400 on the last invoice."))
        db.session.commit()
        m = db.session.get(M.Matter, mid)
        bodies = [n.body for n in update_facts(m)["notes"]]
    assert any("deposition transcript" in b for b in bodies)
    assert not any("owes" in b for b in bodies), "a note about money reached the client-facing facts"
